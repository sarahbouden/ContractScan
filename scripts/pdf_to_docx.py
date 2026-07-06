import os, pdfplumber
from docx import Document

os.makedirs("data/contracts/docx", exist_ok=True)

pdf_dir = "data/contracts/pdf"
for pdf_file in os.listdir(pdf_dir)[:5]:  # 5 contrats en DOCX
    path = os.path.join(pdf_dir, pdf_file)
    with pdfplumber.open(path) as pdf:
        full_text = "\n".join(
            page.extract_text() or "" for page in pdf.pages
        )
    doc = Document()
    doc.add_heading(pdf_file.replace(".pdf", ""), 0)
    for paragraph in full_text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    docx_path = os.path.join(
        "data/contracts/docx",
        pdf_file.replace(".pdf", ".docx")
    )
    doc.save(docx_path)
    print(f"Converti : {docx_path}")
